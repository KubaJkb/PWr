from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(path: str) -> str:
    try:
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"[BŁĄD PDF] Nie można odczytać pliku '{path}': {e}")
        return ""

def extract_text_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        text_parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return '\n'.join(text_parts).strip()
    except Exception as e:
        print(f"[BŁĄD DOCX] Nie można odczytać pliku '{path}': {e}")
        return ""

def load_document(path: str) -> str:
    try:
        if path.lower().endswith(".pdf"):
            return extract_text_from_pdf(path)
        elif path.lower().endswith(".docx"):
            return extract_text_from_docx(path)
        else:
            print(f"[BŁĄD] Nieobsługiwany format pliku: {path}")
            return ""
    except Exception as e:
        print(f"[BŁĄD] Błąd podczas ładowania pliku {path}: {e}")
        return ""
